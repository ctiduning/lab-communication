# Generate QDCTI delivery document (Word format)
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ===== Style setup =====
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Colors
C_PRIMARY = RGBColor(0x1A, 0x52, 0x76)
C_SECONDARY = RGBColor(0x2E, 0x86, 0xC1)
C_LIGHT_BG = RGBColor(0xD6, 0xEA, 0xF8)
C_TEXT = RGBColor(0x2C, 0x3E, 0x50)
C_GREEN = RGBColor(0x27, 0xAE, 0x60)
C_GRAY = RGBColor(0xBD, 0xC3, 0xC7)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_title(text, size=48, color=C_PRIMARY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = C_PRIMARY if level == 1 else C_SECONDARY
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_number(text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if header:
            run.bold = True
            run.font.color.rgb = C_WHITE
            set_cell_shading(cell, '1A5276')
        else:
            run.font.color.rgb = C_TEXT
        if width:
            cell.width = Inches(width)

# ===========================
# COVER PAGE
# ===========================
for _ in range(6):
    doc.add_paragraph()

add_title('QDCTI', 72)
add_title('实验室沟通管理系统', 52)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 30)
run.font.color.rgb = C_GRAY
run.font.size = Pt(14)

add_title('完整代码交付 & 功能使用手册', 28, C_TEXT)
doc.add_paragraph()
add_title('版本 1.0', 18, C_PRIMARY)
add_title('交付日期: 2026年6月11日', 18, C_TEXT)
add_title('前端框架: Vue 3 + Element Plus + Supabase', 14, C_GRAY)

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 本文档包含完整源代码、开发时间线及各角色功能说明 —')
run.font.color.rgb = C_GRAY
run.font.size = Pt(10)
run.italic = True

doc.add_page_break()

# ===========================
# TABLE OF CONTENTS
# ===========================
add_heading('目录', 1)
toc_items = [
    '第一部分：项目概览',
    '  1.1 项目简介',
    '  1.2 技术架构',
    '  1.3 项目目录结构',
    '  1.4 系统角色',
    '第二部分：开发时间线（完整档案）',
    '第三部分：数据库变更SQL',
    '第四部分：角色功能说明 & 使用手册',
    '  管理员（Admin）',
    '  业务端（Business）',
    '  实验室端（Lab）',
    '第五部分：完整源代码清单',
]
for item in toc_items:
    add_body(item)

doc.add_page_break()

# ===========================
# PART 1: PROJECT OVERVIEW
# ===========================
add_heading('第一部分：项目概览', 1)

add_heading('1.1 项目简介', 2)
add_body('QDCTI 实验室沟通管理系统是一个基于 Vue 3 + Element Plus + Supabase 构建的企业级沟通平台，专为实验室与业务端之间的高效协作而设计。系统覆盖消息发送、接收、回复、公告管理、用户管理等全链路功能。')
add_body('系统采用单页应用（SPA）架构，部署于 GitHub Pages，后端使用 Supabase（BaaS）提供数据库、认证、存储和实时功能。')

add_heading('1.2 技术架构', 2)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
header_cells = table.rows[0].cells
for i, txt in enumerate(['层次', '技术', '说明']):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_WHITE
    run.font.name = 'Arial'
    set_cell_shading(header_cells[i], '1A5276')

tech_data = [
    ('前端框架', 'Vue 3', 'Composition API + <script setup>'),
    ('UI组件库', 'Element Plus', '企业级桌面UI组件'),
    ('构建工具', 'Vite 5', '快速开发服务器 + 生产构建'),
    ('后端/BaaS', 'Supabase', 'PostgreSQL + Auth + Storage + Realtime'),
    ('路由', 'Vue Router', 'Hash模式（GitHub Pages适配）'),
    ('部署', 'GitHub Pages', '自动构建部署'),
]
for row_data in tech_data:
    add_table_row(table, [(row_data[0], 1.5), (row_data[1], 1.5), (row_data[2], 3.5)])

add_heading('1.3 项目目录结构', 2)
add_bullet('src/api/              - API层（Supabase数据库操作）')
add_bullet('src/pages/            - 页面组件（共15个页面）')
add_bullet('src/utils/            - 工具函数（Supabase客户端等）')
add_bullet('src/router/           - 路由配置')
add_bullet('src/App.vue           - 根组件')
add_bullet('src/main.js           - 入口文件')
add_bullet('package.json          - 依赖配置')
add_bullet('vite.config.js        - Vite构建配置')
add_bullet('supabase-changes.sql  - 数据库变更SQL')

add_heading('1.4 系统角色', 2)
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
headers = table.rows[0].cells
for i, txt in enumerate(['角色', '页面权限', '核心功能', '路由前缀']):
    headers[i].text = ''
    p = headers[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = C_WHITE
    run.font.name = 'Arial'
    set_cell_shading(headers[i], '1A5276')

role_data = [
    ('管理员', '全部页面', '用户管理/公告/日志/角色切换', '/admin'),
    ('业务端', '业务相关页面', '发起业务沟通/接收实验室消息', '/business-*'),
    ('实验室端', '实验室相关页面', '发起实验室沟通/接收业务消息', '/lab-*'),
]
for row_data in role_data:
    add_table_row(table, [(row_data[0], 1), (row_data[1], 1.5), (row_data[2], 3), (row_data[3], 1.5)])

doc.add_page_break()

# ===========================
# PART 2: DEVELOPMENT TIMELINE
# ===========================
add_heading('第二部分：开发时间线（完整档案）', 1)
add_body('以下为项目从创建到当前的全部开发记录，按时间顺序排列。')

timeline_data = [
    ['2026-06-05 11:44', '系统初始化', '实验室沟通系统初始版本搭建'],
    ['2026-06-05 13:22', '部署配置', '适配 GitHub Pages 部署'],
    ['2026-06-05 14:02', '公告管理', '管理员批量选择删除公告'],
    ['2026-06-05 15:10', '组织架构', '组织架构分栏 + 接收人分组搜索'],
    ['2026-06-05 15:24', '搜索优化', '接收人选择器支持拼音搜索'],
    ['2026-06-05 16:08', '登录修复', '修复登录串号问题'],
    ['2026-06-06 12:11', 'Excel导入', 'Excel拖拽上传批量注册用户'],
    ['2026-06-06 13:07', 'UI美化', '优化界面UI样式设计'],
    ['2026-06-06 14:33', '通讯录', '多选 + 快捷沟通 + 表格优化'],
    ['2026-06-06 15:36', '角色切换', '管理员角色切换功能'],
    ['2026-06-06 19:56', '三级体系', '人员架构三级体系重构'],
    ['2026-06-06 20:55', '公告已读表', '新增 announcement_reads 表'],
    ['2026-06-06 21:28', '附件下载', '消息详情附件列表 + 下载'],
    ['2026-06-06 21:40', '移动适配', '汉堡菜单 + 响应式布局'],
    ['2026-06-06 21:50', '实时订阅', 'Realtime 实时刷新'],
    ['2026-06-07 06:13', '撤回功能', '5分钟内消息撤回'],
    ['2026-06-07 15:19', '点赞系统', '公告/回复点赞点踩'],
    ['2026-06-08 16:55', '四项功能', '草稿暂存/红旗/模糊搜索/追加回复'],
    ['2026-06-11 01:39', '回退重构', '回退到V1重新实现三项需求'],
    ['2026-06-11 05:40', '模式重写', '按接收消息模式重写公告标记'],
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
headers = table.rows[0].cells
for i, txt in enumerate(['日期', '类型', '变更描述']):
    headers[i].text = ''
    p = headers[i].paragraphs[0]
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_WHITE
    run.font.name = 'Arial'
    set_cell_shading(headers[i], '1A5276')

for row_data in timeline_data:
    add_table_row(table, [(row_data[0], 1.8), (row_data[1], 1.2), (row_data[2], 4)])

doc.add_page_break()

# ===========================
# PART 3: DATABASE SQL
# ===========================
add_heading('第三部分：数据库变更SQL', 1)
add_body('请在 Supabase SQL Editor 中执行以下SQL：')
sql = '''-- 1. 公告红旗标记 - announcement_reads 表加字段
ALTER TABLE announcement_reads ADD COLUMN IF NOT EXISTS is_flagged BOOLEAN DEFAULT false;

-- 2. Thread回复 - replies 表加 target_recipient_id
ALTER TABLE replies ADD COLUMN IF NOT EXISTS target_recipient_id UUID REFERENCES auth.users(id) ON DELETE SET NULL;'''
p = doc.add_paragraph()
run = p.add_run(sql)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ===========================
# PART 4: ROLE FEATURE DOCS
# ===========================
add_heading('第四部分：角色功能说明 & 使用手册', 1)

# Helper: add module card
def add_module_card(name, path, features, instructions):
    # Title row
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(f'📌 {name}')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_WHITE
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_cell_shading(cell, '1A5276')

    # Path row
    row = table.add_row()
    cell = row.cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(f'📁 源文件: {path}')
    run.font.size = Pt(9)
    run.font.color.rgb = C_SECONDARY
    run.italic = True
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_cell_shading(cell, 'EBF5FB')

    # Features
    row = table.add_row()
    cell = row.cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run('功能说明：')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    for f in features:
        bp = cell.add_paragraph(f, style='List Bullet')
        for r in bp.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)

    # Instructions
    row = table.add_row()
    cell = row.cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run('使用说明：')
    run.bold = True
    run.font.color.rgb = C_GREEN
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    for i, instr in enumerate(instructions):
        bp = cell.add_paragraph(f'{i+1}. {instr}')
        for r in bp.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)

    doc.add_paragraph()  # spacing

# ---- ADMIN ----
add_heading('管理员（Admin）', 2)
add_body('管理员拥有系统最高权限，可管理用户、发布公告、查看操作日志，并可通过角色切换预览不同视图。')

admin_modules = [
    ('仪表盘（Home）', 'Home.vue',
     ['显示系统概览统计信息', '管理员可切换业务端/实验室端视图', '实时数据刷新'],
     ['登录后默认进入仪表盘', '顶部导航栏可切换不同模块']),
    ('用户管理（Admin）', 'Admin.vue',
     ['用户注册：创建新用户', '批量导入：Excel拖拽上传批量创建', '用户编辑、重置密码', '用户禁用/启用'],
     ['【注册用户】填写表单后提交', '【批量导入】点击Excel导入按钮', '【编辑用户】点击用户行右侧编辑按钮']),
    ('公告管理（Announcements）', 'Announcements.vue',
     ['发布、编辑、删除公告', '已读状态追踪', '每个人可对重要公告做红旗标记', '点赞/点踩反馈', '模糊搜索'],
     ['【发布公告】点击"+发布通知"按钮', '【红旗标记】每行右侧操作栏点击按钮', '【筛选】勾选"仅显示🚩标记"', '【详情】点击行任意位置']),
    ('通知管理（Notifications）', 'Notifications.vue',
     ['系统通知列表查看', '全部标记为已读'],
     ['查看通知列表，点击"全部已读"']),
    ('操作日志（History）', 'History.vue',
     ['管理员操作日志', '按时间过滤查看', '导出Excel'],
     ['进入操作日志页面', '选择日期范围查看', '点击导出']),
    ('通讯录（Organization）', 'Organization.vue',
     ['按部门层级展示用户', '多选用户快捷沟通', '模糊搜索和拼音搜索'],
     ['选择用户后点击"快捷沟通"跳转']),
    ('个人设置（Profile）', 'Profile.vue',
     ['修改个人信息', '修改密码'],
     ['进入个人设置维护信息']),
]
for name, path, features, instructions in admin_modules:
    add_module_card(name, path, features, instructions)

doc.add_page_break()

# ---- BUSINESS ----
add_heading('业务端（Business）', 2)
add_body('业务端用户可向实验室发起沟通请求，查看实验室回复，对重要消息进行红旗标记。')

business_modules = [
    ('发起沟通（BusinessInitiate）', 'BusinessInitiate.vue',
     ['选择沟通类型（付费加急/免费加急/数据质疑/跟单等）', '填写客户/样品信息', '选择接收人（支持多人+部门名片）', '上传附件', '草稿暂存', '编辑重发已撤回消息'],
     ['选择类型→填写信息→选择接收人→发送', '中途可暂存为草稿', '已撤回消息可编辑重发']),
    ('接收消息（BusinessReceive）', 'BusinessReceive.vue',
     ['标签页分类：待处理/已处理/已完结/已撤回/全部', '红旗标记', '快捷回复（同意/拒绝/等我确认后回复）', '关键词搜索（含回复内容）', '点击行任意位置查看详情', '附件下载'],
     ['【查看消息】点击表格任意位置打开详情', '【快捷回复】列表行点击同意/拒绝按钮', '【回复】详情弹窗底部输入框填写', '【红旗标记】行上按钮标记/取消', '【附件下载】详情中点击附件链接']),
    ('已发送消息（SentMessages）', 'SentMessages.vue',
     ['查看已发送消息', '按状态筛选（未回复/已回复/已完结/已撤回）', '标记整体完结', '追加回复', '撤回消息（5分钟内）', '编辑重发', '模糊搜索'],
     ['【完结】详情弹窗点击"标记整体完结"', '【追加回复】列表右侧点击按钮', '【撤回】详情中点击"撤回消息"（限5分钟）']),
    ('通讯录（Organization）', 'Organization.vue', ['查看所有用户', '按部门层级展示', '搜索'], ['查看团队通讯录']),
    ('个人设置（Profile）', 'Profile.vue', ['修改个人信息', '修改密码'], ['维护个人信息']),
]
for name, path, features, instructions in business_modules:
    add_module_card(name, path, features, instructions)

doc.add_page_break()

# ---- LAB ----
add_heading('实验室端（Lab）', 2)
add_body('实验室端用户可接收业务端的沟通请求并回复，也可主动发起沟通。功能与业务端对称。')

lab_modules = [
    ('发起沟通（LabInitiate）', 'LabInitiate.vue',
     ['选择沟通类型', '填写沟通内容及样品信息', '选择接收人', '上传附件', '草稿暂存'],
     ['与业务端发起沟通功能一致', '实验室端专用入口']),
    ('接收消息（LabReceive）', 'LabReceive.vue',
     ['标签页分类', '红旗标记', '快捷回复', '关键词搜索', '点击行查看详情', '附件下载'],
     ['【查看】点击表格任意位置', '【快捷回复】列表行点击按钮', '【红旗】行上按钮标记/取消', '【附件】详情中下载']),
    ('已发送消息（SentMessages）', 'SentMessages.vue',
     ['查看已发送消息', '状态筛选', '标记完结', '追加回复', '撤回消息（5分钟内）', '编辑重发'],
     ['【筛选】顶部标签切换', '【完结】详情中标记整体完结', '【撤回】详情中点击撤回消息']),
    ('通讯录（Organization）', 'Organization.vue', ['查看所有用户', '按部门层级展示', '搜索'], ['查看团队通讯录']),
    ('个人设置（Profile）', 'Profile.vue', ['修改个人信息', '修改密码'], ['维护个人信息']),
]
for name, path, features, instructions in lab_modules:
    add_module_card(name, path, features, instructions)

doc.add_page_break()

# ===========================
# PART 5: SOURCE CODE LIST
# ===========================
add_heading('第五部分：完整源代码清单', 1)
add_body('以下为项目所有源代码文件的完整路径清单，完整源码已同步存放于桌面「QDCTI交付资料」文件夹中。')

add_heading('核心文件', 2)
add_bullet('src/main.js              - 应用入口')
add_bullet('src/App.vue              - 根组件（导航布局）')
add_bullet('src/api/index.js          - API层（所有Supabase数据库操作）')
add_bullet('src/utils/supabase.js     - Supabase客户端配置')
add_bullet('src/router/index.js       - 路由配置')
add_bullet('package.json              - 依赖配置文件')
add_bullet('vite.config.js            - Vite构建配置')
add_bullet('supabase-changes.sql      - 数据库变更SQL')

add_heading('页面组件（src/pages/）', 2)
pages = [
    'Login.vue', 'Home.vue', 'Admin.vue', 'Announcements.vue',
    'Notifications.vue', 'BusinessInitiate.vue', 'BusinessReceive.vue',
    'LabInitiate.vue', 'LabReceive.vue', 'SentMessages.vue',
    'Organization.vue', 'Profile.vue', 'History.vue',
    'Communications.vue', 'Orders.vue'
]
for p in pages:
    add_bullet(p)

add_heading('API层核心方法', 2)
add_bullet('userAPI - 用户查询（getByRole / getAll）')
add_bullet('authAPI - 认证注册（register / login）')
add_bullet('communicationAPI - 消息CRUD（create / getAll / getById / createReply / recall等）')
add_bullet('announcementAPI - 公告管理（create / list / update / delete / toggleAnnouncementFlag / markAsRead）')
add_bullet('reactionAPI - 点赞点踩（toggle / getStatsBatch / getDetail）')
add_bullet('storageAPI - 文件上传')
add_bullet('uploadAPI - 上传工具')
add_bullet('logAPI - 操作日志')

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')
run = p.add_run('— 文档结束 —')
run.font.size = Pt(14)
run.font.color.rgb = C_GRAY
run.italic = True

# Save
output_path = 'C:/Users/大腚使者/Desktop/QDCTI交付资料/QDCTI_完整交付文档.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
