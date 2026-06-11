# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ===== 样式 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

C_PRIMARY = RGBColor(0x1A, 0x52, 0x76)
C_SECONDARY = RGBColor(0x2E, 0x86, 0xC1)
C_TEXT = RGBColor(0x2C, 0x3E, 0x50)
C_GREEN = RGBColor(0x27, 0xAE, 0x60)
C_RED = RGBColor(0xE7, 0x4C, 0x3C)
C_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
C_GRAY = RGBColor(0x95, 0xA5, 0xA6)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_BG = RGBColor(0xD6, 0xEA, 0xF8)
C_WARN_BG = RGBColor(0xFD, 0xED, 0xED)

def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_title(text, size=48, color=C_PRIMARY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = C_PRIMARY if level == 1 else C_SECONDARY
        r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Arial'; r.font.size = Pt(11)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bold_body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True; r.font.size = Pt(11)
        r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    else:
        p.clear()
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_number(text):
    p = doc.add_paragraph(text, style='List Number')
    for r in p.runs:
        r.font.name = 'Arial'; r.font.size = Pt(11)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_warning(text):
    p = doc.add_paragraph()
    r = p.add_run('⚠️ ' + text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_RED
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_tip(text):
    p = doc.add_paragraph()
    r = p.add_run('💡 ' + text)
    r.font.size = Pt(10); r.font.color.rgb = C_SECONDARY
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_step(num, text):
    p = doc.add_paragraph()
    r = p.add_run(f'步骤{num}：')
    r.bold = True; r.font.size = Pt(11)
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()
add_title('QDCTI 实验室沟通管理系统', 48)
add_title('使 用 说 明 书', 36, C_SECONDARY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('━' * 30)
r.font.color.rgb = C_GRAY; r.font.size = Pt(14)
add_title('版本 1.0 ｜ 2026年6月', 18, C_TEXT)
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('本文档仅包含功能使用说明，不涉及技术实现细节')
r.font.size = Pt(11); r.font.color.rgb = C_GRAY; r.italic = True
doc.add_page_break()

# ===== 目录 =====
add_heading('目 录', 1)
toc = [
    '一、重要提示（必读）',
    '二、系统概述',
    '三、管理员模块说明',
    '  3.1 仪表盘',
    '  3.2 用户管理',
    '  3.3 通知公告',
    '  3.4 操作日志',
    '  3.5 通讯录',
    '  3.6 个人设置',
    '四、业务端模块说明',
    '  4.1 发起沟通',
    '  4.2 接收消息',
    '  4.3 已发送消息',
    '  4.4 通讯录',
    '  4.5 个人设置',
    '五、实验室端模块说明',
    '  5.1 发起沟通',
    '  5.2 接收消息',
    '  5.3 已发送消息',
    '  5.4 通讯录',
    '  5.5 个人设置',
    '六、常见问题与解决方法',
    '七、获取帮助',
]
for t in toc:
    add_body(t)
doc.add_page_break()

# ===== 一、重要提示 =====
add_heading('一、重要提示（必读）', 1)

# 重点标注 Ctrl+F5
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('⚠️ ⚠️ ⚠️ 每次登录前请务必执行以下操作 ⚠️ ⚠️ ⚠️')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = C_RED
r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ctrl + F5  强制刷新页面')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = C_RED
r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（键盘左下角的 Ctrl 键 + 键盘顶部的 F5 功能键同时按下）')
r.font.size = Pt(11); r.font.color.rgb = C_TEXT
r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

add_body('为什么要这样做？')
add_bullet('系统更新后，浏览器可能缓存了旧版本的页面文件（JavaScript、CSS等）')
add_bullet('使用旧缓存会导致新功能不显示、按钮点不动、样式错乱等问题')
add_bullet('Ctrl+F5 会强制浏览器从服务器重新下载所有文件，而不是使用本地缓存')
add_tip('每次打开系统时，养成习惯先按 Ctrl+F5，再输入账号密码登录。')
doc.add_paragraph()

add_body('其他注意事项：')
add_bullet('推荐使用 Chrome（谷歌浏览器）或 Edge 浏览器访问')
add_bullet('屏幕分辨率建议 1366×768 以上获得最佳体验')

# ===== 二、系统概述 =====
doc.add_page_break()
add_heading('二、系统概述', 1)
add_body('QDCTI 实验室沟通管理系统是一个专为实验室与业务端之间高效协作而设计的在线沟通平台。系统支持消息发送、接收、回复、公告发布、用户管理等功能，覆盖实验室日常工作的全链路沟通需求。')

add_heading('2.1 系统角色', 2)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
headers = table.rows[0].cells
for i, txt in enumerate(['角色', '职责', '可访问的模块']):
    headers[i].text = ''
    p = headers[i].paragraphs[0]
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_WHITE
    r.font.name = 'Arial'
    set_shading(headers[i], '1A5276')

roles = [
    ('管理员', '系统运维管理', '全部模块 + 用户管理 + 公告管理 + 日志'),
    ('业务端', '向实验室发起沟通请求', '发起沟通/接收消息/已发送/通讯录/个人设置'),
    ('实验室端', '接收并回复业务端的消息', '发起沟通/接收消息/已发送/通讯录/个人设置'),
]
for role, duty, modules in roles:
    row = table.add_row()
    for i, txt in enumerate([role, duty, modules]):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10); r.font.name = 'Arial'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_heading('2.2 访问方式', 2)
add_body('在浏览器地址栏中输入系统网址，按回车即可打开登录页面。输入管理员分配的用户名和密码进行登录。')

# ===== 三、管理员 =====
doc.add_page_break()
add_heading('三、管理员模块说明', 1)
add_body('管理员拥有系统的最高权限，可以管理用户、发布公告、查看系统日志，还可以通过角色切换功能预览业务端和实验室端的界面。')

# 3.1 仪表盘
add_heading('3.1 仪表盘', 2)
add_bold_body('功能简介：')
add_body('仪表盘是登录后默认进入的页面，展示系统运行概况，包括待处理消息数量、未读公告数量等统计信息。')
add_bold_body('操作说明：')
add_step('1', '登录后自动进入仪表盘页面')
add_step('2', '顶部导航栏可点击切换至其他功能模块')
add_step('3', '管理员可在顶部点击角色切换按钮，一键切换到业务端或实验室端视图')
add_tip('角色切换后只能查看对应角色的界面，但不会影响您的管理员身份，切换回来即可恢复。')

# 3.2 用户管理
add_heading('3.2 用户管理（Admin）', 2)
add_bold_body('功能简介：')
add_body('用户管理是管理员的核心模块，用于创建、编辑、禁用/启用系统用户，支持单个添加和批量导入两种方式。')
add_bold_body('操作说明：')
add_step('1', '进入管理员控制台页面')
add_step('2', '【单个注册】填写用户名、密码、姓名、角色、部门等信息，点击注册')
add_step('3', '【批量导入】准备好 Excel 文件（格式可在页面下载模板），点击"Excel导入"上传')
add_step('4', '【编辑用户】点击用户列表中的编辑按钮，修改信息或重置密码')
add_step('5', '【禁用/启用】点击用户状态开关，禁用后该用户无法登录系统')
add_tip('批量导入时请确保 Excel 文件格式与模板一致，否则可能导入失败。')
add_tip('重置密码后，建议让用户登录后自行修改密码。')

# 3.3 通知公告
add_heading('3.3 通知公告（Announcements）', 2)
add_bold_body('功能简介：')
add_body('通知公告模块用于向系统内用户发布公告通知，支持选择发送范围，并跟踪用户的已读状态。')
add_bold_body('操作说明：')
add_step('1', '【发布公告】点击"+发布通知"按钮，填写标题、内容和发送范围后发布')
add_step('2', '【查看公告】点击公告表格的任意位置打开详情查看')
add_step('3', '【红旗标记】每条公告右侧有🚩按钮，点击后可标记为重要，方便后续快速筛选')
add_step('4', '【筛选】勾选"仅显示🚩标记"可只查看自己标记了红旗的公告')
add_step('5', '【点赞反馈】打开公告详情后，底部可点击👍赞或👎踩')
add_step('6', '【编辑/删除】管理员可在公告列表右侧点击编辑或删除按钮')
add_tip('公告发布后不可修改接收范围，请发布前确认。')

# 3.4 操作日志
add_heading('3.4 操作日志（History）', 2)
add_bold_body('功能简介：')
add_body('操作日志记录管理员在系统中的所有关键操作，包括用户注册、禁用、密码重置等，支持按时间范围查询和导出。')
add_bold_body('操作说明：')
add_step('1', '进入历史记录页面')
add_step('2', '可选择日期范围筛选日志')
add_step('3', '点击导出按钮可将日志导出为 Excel 文件')

# 3.5 通讯录
add_heading('3.5 通讯录（Organization）', 2)
add_bold_body('功能简介：')
add_body('通讯录按部门层级展示所有系统用户的信息，支持按姓名、拼音、部门进行搜索。')
add_bold_body('操作说明：')
add_step('1', '进入通讯录页面')
add_step('2', '在搜索框输入姓名或部门名称搜索')
add_step('3', '选中用户后点击"快捷沟通"自动跳转到发起沟通页面')
add_tip('支持拼音首字母搜索，例如输入"zs"可搜索到"张三"。')

# 3.6 个人设置
add_heading('3.6 个人设置（Profile）', 2)
add_bold_body('功能简介：')
add_body('用于修改登录用户的个人信息和登录密码。')
add_bold_body('操作说明：')
add_step('1', '进入个人设置页面')
add_step('2', '修改姓名、电话、邮箱等信息后点击保存')
add_step('3', '如需修改密码，输入当前密码和新密码后确认')

# ===== 四、业务端 =====
doc.add_page_break()
add_heading('四、业务端模块说明（Business）', 1)
add_body('业务端用户可向实验室发起沟通请求，查看实验室的回复，对重要消息进行标记。')

# 4.1 发起沟通
add_heading('4.1 发起沟通（BusinessInitiate）', 2)
add_bold_body('功能简介：')
add_body('向实验室发送沟通请求。支持选择沟通类型、填写客户和样品信息、选择多个接收人、上传附件等。')
add_bold_body('操作说明：')
add_step('1', '进入发起沟通页面')
add_step('2', '选择沟通类型（付费加急/免费加急/数据质疑/跟单/咨询/不合格沟通/数据确认等）')
add_step('3', '填写客户名称、样品短号、测试项目等详细信息')
add_step('4', '选择接收人（支持搜索和拼音查找）')
add_step('5', '如有需要可上传附件图片')
add_step('6', '填写完成后点击"发送"按钮')
add_tip('填写中途可点击"保存草稿"，稍后继续编辑。')
add_tip('已撤回的消息可在已发送页面点击"编辑重发"一键恢复。')

# 4.2 接收消息
add_heading('4.2 接收消息（BusinessReceive）', 2)
add_bold_body('功能简介：')
add_body('查看来自实验室的消息，支持快捷回复、红旗标记、关键字搜索等。')
add_bold_body('标签页说明：')
add_bullet('待处理：', '显示尚未回复的消息')
add_bullet('已处理：', '显示已回复但未完结的消息')
add_bullet('已完结：', '显示发起人已标记完结的消息')
add_bullet('已被撤回：', '显示发送方已撤回的消息')
add_bullet('全部：', '显示所有消息')
add_bold_body('操作说明：')
add_step('1', '点击表格任意位置打开消息详情')
add_step('2', '【快捷回复】在列表中直接点击"同意/拒绝/等我确认"按钮快速回复')
add_step('3', '【手动回复】在详情弹窗底部输入框填写回复内容，点击"发送回复"')
add_step('4', '【红旗标记】点击行上的🚩按钮标记重要消息（仅自己可见）')
add_step('5', '【附件下载】在详情中点击附件链接即可下载')
add_step('6', '【搜索】在搜索框输入关键词，可搜索消息内容和回复记录')
add_tip('点击"等我确认后回复"不会标记为已回复，仅作为占位提醒。')

# 4.3 已发送消息
add_heading('4.3 已发送消息（SentMessages）', 2)
add_bold_body('功能简介：')
add_body('查看自己发出的所有沟通消息，追踪各接收人的回复状态。')
add_bold_body('状态标签说明：')
add_bullet('未回复：', '所有接收人都还没有回复')
add_bullet('已回复：', '至少有一个接收人已回复')
add_bullet('已完结：', '发起人主动标记了整体完结')
add_bullet('已撤回：', '已撤回的消息（5分钟内可撤回）')
add_bold_body('操作说明：')
add_step('1', '按顶部标签页筛选不同状态的消息')
add_step('2', '【查看详情】点击任意位置打开详情')
add_step('3', '【完结沟通】确认沟通已完成后，在详情弹窗点击"标记整体完结"')
add_step('4', '【追加回复】对已回复的沟通可补充回复内容，可选择"回复全部"或"仅回复某人"')
add_step('5', '【撤回消息】如有误发，在发送后5分钟内可撤回（需填写撤回原因）')
add_step('6', '【编辑重发】撤回后可在"已撤回"标签页点击"编辑重发"重新发送')
add_tip('撤回时限为发送后5分钟，超时将无法撤回。')

# 4.4 4.5
add_heading('4.4 通讯录', 2)
add_body('查看系统内所有用户的联系信息，按部门层级展示，支持搜索。功能与管理员通讯录一致。')

add_heading('4.5 个人设置', 2)
add_body('修改个人信息和密码。')

# ===== 五、实验室端 =====
doc.add_page_break()
add_heading('五、实验室端模块说明（Lab）', 1)
add_body('实验室端用户可接收业务端的沟通请求并回复，也可主动发起沟通。功能与业务端对称。')

add_heading('5.1 发起沟通（LabInitiate）', 2)
add_bold_body('功能简介：')
add_body('向业务端发送沟通请求。功能与业务端发起沟通一致，为实验室端专用入口。')
add_bold_body('操作说明：')
add_step('1', '进入发起沟通页面（实验室端入口）')
add_step('2', '选择沟通类型并填写相关信息')
add_step('3', '选择接收人（支持搜索和拼音查找）')
add_step('4', '上传附件（如有需要）')
add_step('5', '点击发送')
add_tip('支持草稿暂存和编辑重发功能。')

add_heading('5.2 接收消息（LabReceive）', 2)
add_bold_body('功能简介：')
add_body('查看来自业务端的消息。标签页和操作方式与业务端接收消息一致。')
add_bold_body('操作说明：')
add_step('1', '点击表格任意位置打开消息详情')
add_step('2', '在列表中直接点击"同意/拒绝/等我确认"按钮快捷回复')
add_step('3', '点击🚩按钮标记重要消息')
add_step('4', '在搜索框输入关键词搜索')
add_step('5', '在详情中点击附件链接下载文件')
add_tip('回复内容支持输入自定义文本，不仅限于快捷回复。')

add_heading('5.3 已发送消息（SentMessages）', 2)
add_body('功能与操作同业务端已发送消息一致。可查看发送历史、追踪回复状态、撤回消息、编辑重发。')

add_heading('5.4 通讯录', 2)
add_body('查看所有用户信息，按部门展示，支持搜索。')

add_heading('5.5 个人设置', 2)
add_body('修改个人信息和密码。')

# ===== 六、常见问题 =====
doc.add_page_break()
add_heading('六、常见问题与解决方法', 1)

faqs = [
    ('Q：打开页面显示空白或样式错乱？',
     'A：按 Ctrl+F5 强制刷新页面，清除浏览器缓存后重试。',
     C_RED),
    ('Q：点击按钮没反应？',
     'A：先按 Ctrl+F5 刷新，如仍不行请检查是否网络连接正常。',
     C_RED),
    ('Q：登录提示"账号或密码错误"？',
     'A：请联系管理员重置密码。如确认密码正确，请按 Ctrl+F5 刷新后重试。',
     C_ORANGE),
    ('Q：发布公告后发现内容写错了？',
     'A：管理员可以在公告列表点击"编辑"按钮修改已发布的公告。',
     C_ORANGE),
    ('Q：消息发错了怎么办？',
     'A：发送后5分钟内可以在已发送消息页面点击"撤回消息"，填写原因后撤回。',
     C_ORANGE),
    ('Q：如何知道对方已经看到我的消息？',
     'A：在消息详情中可以看到每个接收人的已读/未读状态。',
     C_SECONDARY),
    ('Q：什么是"等我确认后回复"？',
     'A：这是一个占位回复，表示您已看到消息但需要时间确认，不会标记为已回复。',
     C_SECONDARY),
    ('Q：如何只查看自己标记了红旗的消息？',
     'A：在接收消息和公告页面均有"仅显示🚩标记"的筛选复选框，勾选即可。',
     C_SECONDARY),
    ('Q：为什么我看不到某些页面？',
     'A：不同角色有不同的页面权限。如果确认权限不足，请联系管理员。',
     C_SECONDARY),
]

for question, answer, color in faqs:
    p = doc.add_paragraph()
    r = p.add_run(question)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = color
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p = doc.add_paragraph()
    r = p.add_run(answer)
    r.font.size = Pt(11)
    r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ===== 七、获取帮助 =====
add_heading('七、获取帮助', 1)
add_body('如果您在使用过程中遇到任何问题，可以通过以下方式获取帮助：')
add_bullet('联系系统管理员')
add_bullet('查阅本文档的相关章节')
add_bullet('按 Ctrl+F5 强制刷新后重试（超过半数的问题通过此方法即可解决）')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 使用说明书结束 —')
r.font.size = Pt(14); r.font.color.rgb = C_GRAY; r.italic = True
r.font.name = 'Arial'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 保存 =====
output = 'C:/Users/大腚使者/Desktop/QDCTI交付资料/QDCTI_使用说明书.docx'
doc.save(output)
print(f'✅ 已生成: {output}')
